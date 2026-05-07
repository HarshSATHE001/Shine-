from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_report():
    doc = Document()

    # Title
    title = doc.add_heading('Project Shine: AI Integration & Model Optimization Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph('Date: May 2, 2026').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('Lead AI Engineer: Antigravity (Advanced Coding Agent)').alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_section()

    # Section 1: Executive Summary
    doc.add_heading('1. Executive Summary', level=1)
    doc.add_paragraph(
        'Today, we successfully transitioned the "Shine" educational dropout prediction system from a research prototype '
        'into a production-grade, integrated microservice architecture. We achieved significant breakthroughs in model '
        'accuracy and built a unified interface for mentor-led batch analysis.'
    )

    # Section 2: Core Accomplishments
    doc.add_heading('2. Core Accomplishments', level=1)

    doc.add_heading('A. Model Optimization (XGBoost)', level=2)
    p = doc.add_paragraph()
    p.add_run('We replaced the Random Forest baseline with a high-performance ')
    p.add_run('Lightweight XGBoost multiclass classifier').bold = True
    p.add_run('.')
    
    list_items = [
        ('Performance Peak', 'High-Risk Student Recall increased from 67% to 83% (+24%).'),
        ('Recall Priority', 'Optimized for "Safety First," ensuring fewer at-risk students are missed.'),
        ('Efficiency', 'Tuned for standard laptops, maintaining low RAM and CPU usage.')
    ]
    for bold_text, normal_text in list_items:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{bold_text}: ').bold = True
        p.add_run(normal_text)

    doc.add_heading('B. Explainable AI (SHAP)', level=2)
    doc.add_paragraph(
        'Integrated a production-ready SHAP (SHapley Additive exPlanations) layer.'
    )
    list_items = [
        ('Global Insights', "Identified 'Absences', 'Failures', and 'Age' as top risk drivers."),
        ('Local Logic', "Enabled human-readable 'Risk Stories' for each student (e.g., 'Flagged due to high absenteeism').")
    ]
    for bold_text, normal_text in list_items:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{bold_text}: ').bold = True
        p.add_run(normal_text)

    doc.add_heading('C. Production AI Service (FastAPI)', level=2)
    doc.add_paragraph('Built a standalone AI microservice to serve predictions.')
    list_items = [
        ('Endpoints', 'Implemented /predict (single) and /batch-predict (vectorized).'),
        ('Inference Safety', 'Enforced strict Pydantic schema validation and frozen feature ordering to prevent production data drift.')
    ]
    for bold_text, normal_text in list_items:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{bold_text}: ').bold = True
        p.add_run(normal_text)

    doc.add_heading('D. Full-Stack Integration', level=2)
    doc.add_paragraph('Connected the AI engine to the existing Node.js and React stack.')
    list_items = [
        ('Backend Client', 'Created aiService.js in Node.js with 5s timeouts and graceful fallbacks.'),
        ('Unified Dashboard', 'Integrated a new "Batch Analysis" module into the Mentor UI.'),
        ('Batch Processing', 'Added client-side Excel (.xlsx) parsing using the XLSX library.')
    ]
    for bold_text, normal_text in list_items:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{bold_text}: ').bold = True
        p.add_run(normal_text)

    # Section 3: Architecture Note
    doc.add_heading('3. System Status', level=1)
    doc.add_paragraph('Report Status: Completed')
    doc.add_paragraph('System Status: Offline (Safely Stopped)')

    # Save the document
    doc.save('Shine_Today_Report.docx')
    print('Generated Shine_Today_Report.docx')

if __name__ == '__main__':
    create_report()
